import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import io
import base64
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
import json
from datetime import datetime
import numpy as np

# === –§–£–ù–ö–¶–ò–ò –î–õ–Ø –†–ê–ë–û–¢–´ –° –ß–ï–†–ù–û–í–ò–ö–ê–ú–ò ===
def save_draft(data, module_data_list, defects_df):
    """–°–æ—Ö—Ä–∞–Ω—è–µ—Ç –¥–∞–Ω–Ω—ã–µ —Ñ–æ—Ä–º—ã –≤ —Å—Ç—Ä—É–∫—Ç—É—Ä—É –¥–ª—è —á–µ—Ä–Ω–æ–≤–∏–∫–∞"""
    draft = {
        "saved_at": datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
        "data": data,
        "modules": [
            {
                "title": m["title"],
                "df": m["df"].to_dict(orient="records") if not m["df"].empty else []
            }
            for m in module_data_list
        ],
        "defects": defects_df.to_dict(orient="records") if not defects_df.empty else []
    }
    return json.dumps(draft, ensure_ascii=False, indent=2)

def load_draft(json_content):
    """–í–æ—Å—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ—Ç –¥–∞–Ω–Ω—ã–µ –∏–∑ —á–µ—Ä–Ω–æ–≤–∏–∫–∞"""
    try:
        draft = json.loads(json_content)
        
        # –í–æ—Å—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –¥–∞–Ω–Ω—ã–µ —Ñ–æ—Ä–º—ã
        data = draft.get("data", {})
        
        # –í–æ—Å—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –¥–µ—Ñ–µ–∫—Ç—ã
        defects_records = draft.get("defects", [])
        defects_df = pd.DataFrame(
            defects_records,
            columns=["ID", "–ú–æ–¥—É–ª—å", "–ó–∞–≥–æ–ª–æ–≤–æ–∫", "–°–µ—Ä—å—ë–∑–Ω–æ—Å—Ç—å", "–°—Ç–∞—Ç—É—Å"]
        ) if defects_records else pd.DataFrame(columns=["ID", "–ú–æ–¥—É–ª—å", "–ó–∞–≥–æ–ª–æ–≤–æ–∫", "–°–µ—Ä—å—ë–∑–Ω–æ—Å—Ç—å", "–°—Ç–∞—Ç—É—Å"])
        
        # –í–æ—Å—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –º–æ–¥—É–ª–∏
        modules = []
        for mod in draft.get("modules", []):
            df_records = mod.get("df", [])
            df = pd.DataFrame(
                df_records,
                columns=["ID", "–°—Ü–µ–Ω–∞—Ä–∏–π", "–°—Ç–∞—Ç—É—Å", "–ö–æ–º–º–µ–Ω—Ç–∞—Ä–∏–π"]
            ) if df_records else pd.DataFrame(columns=["ID", "–°—Ü–µ–Ω–∞—Ä–∏–π", "–°—Ç–∞—Ç—É—Å", "–ö–æ–º–º–µ–Ω—Ç–∞—Ä–∏–π"])
            modules.append({"title": mod["title"], "df": df})
        
        return data, modules, defects_df, draft.get("saved_at", "–Ω–µ–∏–∑–≤–µ—Å—Ç–Ω–æ")
    except Exception as e:
        st.error(f"‚ùå –û—à–∏–±–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ —á–µ—Ä–Ω–æ–≤–∏–∫–∞: {str(e)}")
        return None, None, None, None

# === –í–°–ü–û–ú–û–ì–ê–¢–ï–õ–¨–ù–´–ï –§–£–ù–ö–¶–ò–ò –î–õ–Ø DOCX ===
def set_col_width(table, col_idx, width_twips):
    """–£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ—Ç —à–∏—Ä–∏–Ω—É –∫–æ–ª–æ–Ω–∫–∏ –≤ —Ç–∞–±–ª–∏—Ü–µ DOCX (–≤ —Ç–≤–∏–ø–∞—Ö)"""
    for row in table.rows:
        cell = row.cells[col_idx]
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_w = OxmlElement('w:tcW')
        tc_w.set(qn('w:w'), str(width_twips))
        tc_w.set(qn('w:type'), 'dxa')
        tc_pr.append(tc_w)

def add_table_from_df(doc, df, col_widths=None):
    """–î–æ–±–∞–≤–ª—è–µ—Ç —Ç–∞–±–ª–∏—Ü—É –∏–∑ DataFrame –≤ –¥–æ–∫—É–º–µ–Ω—Ç DOCX"""
    if df.empty or len(df.columns) == 0:
        doc.add_paragraph("–ù–µ—Ç –¥–∞–Ω–Ω—ã—Ö –¥–ª—è –æ—Ç–æ–±—Ä–∞–∂–µ–Ω–∏—è")
        return
    
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    
    # –ó–∞–≥–æ–ª–æ–≤–∫–∏
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = str(column)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Calibri Light'
    
    # –î–∞–Ω–Ω—ã–µ
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text = str(value) if pd.notna(value) else ""
            row_cells[i].text = cell_text
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri Light'
    
    # –£—Å—Ç–∞–Ω–æ–≤–∫–∞ —à–∏—Ä–∏–Ω—ã –∫–æ–ª–æ–Ω–æ–∫ (–µ—Å–ª–∏ –∑–∞–¥–∞–Ω–∞)
    if col_widths:
        for col_idx, width in enumerate(col_widths):
            set_col_width(table, col_idx, width)

# === –ì–ï–ù–ï–†–ê–¶–ò–Ø DOCX –û–¢–ß–Å–¢–ê ===
def generate_docx(data, module_data_list, defects_df):
    doc = Document()
    
    # –°—Ç–∏–ª—å –¥–æ–∫—É–º–µ–Ω—Ç–∞
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri Light'
    font.size = Pt(11)
    
    # –ó–∞–≥–æ–ª–æ–≤–æ–∫
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(data["report_title"])
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = 'Calibri Light'
    doc.add_paragraph()
    
    # –û—Å–Ω–æ–≤–Ω–∞—è –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è (—Ç–∞–±–ª–∏—Ü–∞ 25%/75%)
    info_table = doc.add_table(rows=7, cols=2)
    info_table.style = 'Table Grid'
    
    fields = [
        ("–ü—Ä–æ–µ–∫—Ç", data["project"]),
        ("–¢–∏–ø –ø—Ä–∏–ª–æ–∂–µ–Ω–∏—è", data["app_type"]),
        ("–í–µ—Ä—Å–∏—è", data["version"]),
        ("–ü–µ—Ä–∏–æ–¥ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è", data["test_period"]),
        ("–î–∞—Ç–∞ —Ñ–æ—Ä–º–∏—Ä–æ–≤–∞–Ω–∏—è –æ—Ç—á—ë—Ç–∞", data["report_date"]),
        ("–ò–Ω–∂–µ–Ω–µ—Ä –ø–æ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—é", data["engineer"]),
        ("–°—Ç–∞—Ç—É—Å —Ä–µ–ª–∏–∑–∞", data["release_status"]),
    ]
    
    for i, (label, value) in enumerate(fields):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = str(value)
        # –ñ–∏—Ä–Ω—ã–π —à—Ä–∏—Ñ—Ç –¥–ª—è –ª–µ–π–±–ª–æ–≤
        for paragraph in info_table.cell(i, 0).paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Calibri Light'
                run.font.size = Pt(10)
        # –û–±—ã—á–Ω—ã–π —à—Ä–∏—Ñ—Ç –¥–ª—è –∑–Ω–∞—á–µ–Ω–∏–π
        for paragraph in info_table.cell(i, 1).paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Calibri Light'
                run.font.size = Pt(10)
    
    # –£—Å—Ç–∞–Ω–æ–≤–∫–∞ —à–∏—Ä–∏–Ω—ã –∫–æ–ª–æ–Ω–æ–∫ 25%/75% (25% = 1800 —Ç–≤–∏–ø–æ–≤ –æ—Ç –æ–±—â–µ–π —à–∏—Ä–∏–Ω—ã ~7200)
    set_col_width(info_table, 0, 1800)
    set_col_width(info_table, 1, 5400)
    doc.add_paragraph()
    
    # –ö—Ä–∞—Ç–∫–æ–µ —Ä–µ–∑—é–º–µ
    doc.add_paragraph("–ö—Ä–∞—Ç–∫–æ–µ —Ä–µ–∑—é–º–µ", style='Heading 2')
    
    # –ú–µ—Ç—Ä–∏–∫–∏ (—Ç–∞–±–ª–∏—Ü–∞ 25%/75%)
    metrics_table = doc.add_table(rows=4, cols=2)
    metrics_table.style = 'Table Grid'
    
    metrics = [
        ("–ö–æ–ª–∏—á–µ—Å—Ç–≤–æ –¥–µ—Ñ–µ–∫—Ç–æ–≤ (S1)", f"{data['s1']}"),
        ("–ö–æ–ª–∏—á–µ—Å—Ç–≤–æ –¥–µ—Ñ–µ–∫—Ç–æ–≤ (S2)", f"{data['s2']}"),
        ("–í—Å–µ–≥–æ —Ç–µ—Å—Ç-–∫–µ–π—Å–æ–≤", f"{data['total_tc']}"),
        ("–ü—Ä–æ–π–¥–µ–Ω–æ —É—Å–ø–µ—à–Ω–æ", f"{data['pass']}"),
    ]
    
    for i, (label, value) in enumerate(metrics):
        metrics_table.cell(i, 0).text = label
        metrics_table.cell(i, 1).text = value
        for paragraph in metrics_table.cell(i, 0).paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Calibri Light'
                run.font.size = Pt(10)
        for paragraph in metrics_table.cell(i, 1).paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Calibri Light'
                run.font.size = Pt(10)
    
    set_col_width(metrics_table, 0, 1800)
    set_col_width(metrics_table, 1, 5400)
    doc.add_paragraph()
    
    # –†–∏—Å–∫–∏ –∏ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏
    if data["risk"].strip():
        doc.add_paragraph("–†–∏—Å–∫–∏", style='Heading 3')
        doc.add_paragraph(data["risk"])
    
    if data["recommendation"].strip():
        doc.add_paragraph("–†–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏", style='Heading 3')
        doc.add_paragraph(data["recommendation"])
    doc.add_paragraph()
    
    # –î–∏–∞–≥—Ä–∞–º–º—ã (–∑–∞–≥–ª—É—à–∫–∏ —Å –æ–ø–∏—Å–∞–Ω–∏–µ–º)
    doc.add_paragraph("–î–∏–∞–≥—Ä–∞–º–º—ã", style='Heading 2')
    doc.add_paragraph("–†–∏—Å. 1. –†–∞—Å–ø—Ä–µ–¥–µ–ª–µ–Ω–∏–µ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è")
    doc.add_paragraph("[–î–∏–∞–≥—Ä–∞–º–º–∞ –±—É–¥–µ—Ç –≤—Å—Ç–∞–≤–ª–µ–Ω–∞ –≤—Ä—É—á–Ω—É—é]")
    doc.add_paragraph()
    doc.add_paragraph("–†–∏—Å. 2. –†–∞—Å–ø—Ä–µ–¥–µ–ª–µ–Ω–∏–µ –¥–µ—Ñ–µ–∫—Ç–æ–≤ –ø–æ —Å–µ—Ä—å—ë–∑–Ω–æ—Å—Ç–∏")
    doc.add_paragraph("[–î–∏–∞–≥—Ä–∞–º–º–∞ –±—É–¥–µ—Ç –≤—Å—Ç–∞–≤–ª–µ–Ω–∞ –≤—Ä—É—á–Ω—É—é]")
    doc.add_paragraph()
    
    # –ö–æ–Ω—Ç–µ–∫—Å—Ç —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è
    doc.add_paragraph("–ö–æ–Ω—Ç–µ–∫—Å—Ç —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è", style='Heading 2')
    
    context_table = doc.add_table(rows=5, cols=2)
    context_table.style = 'Table Grid'
    
    context_fields = [
        ("–£—Å—Ç—Ä–æ–π—Å—Ç–≤–æ / –ë—Ä–∞—É–∑–µ—Ä", data["device_browser"]),
        ("–û–° / –ü–ª–∞—Ç—Ñ–æ—Ä–º–∞", data["os_platform"]),
        ("–°–±–æ—Ä–∫–∞", data["build"]),
        ("URL –æ–∫—Ä—É–∂–µ–Ω–∏—è", data["env_url"]),
        ("–ò–Ω—Å—Ç—Ä—É–º–µ–Ω—Ç—ã", data["tools"]),
    ]
    
    for i, (label, value) in enumerate(context_fields):
        context_table.cell(i, 0).text = label
        context_table.cell(i, 1).text = str(value)
        for paragraph in context_table.cell(i, 0).paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Calibri Light'
                run.font.size = Pt(10)
        for paragraph in context_table.cell(i, 1).paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Calibri Light'
                run.font.size = Pt(10)
    
    set_col_width(context_table, 0, 1800)
    set_col_width(context_table, 1, 5400)
    doc.add_paragraph()
    
    if data["methodology"].strip():
        doc.add_paragraph("–ú–µ—Ç–æ–¥–æ–ª–æ–≥–∏—è —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è", style='Heading 3')
        doc.add_paragraph(data["methodology"])
        doc.add_paragraph()
    
    # –†–µ–∑—É–ª—å—Ç–∞—Ç—ã –ø–æ –º–æ–¥—É–ª—è–º
    doc.add_paragraph("–†–µ–∑—É–ª—å—Ç–∞—Ç—ã —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è –ø–æ –º–æ–¥—É–ª—è–º", style='Heading 2')
    
    for module in module_data_list:
        doc.add_paragraph(module["title"], style='Heading 3')
        add_table_from_df(doc, module["df"], col_widths=[1800, 5400, 1800, 5400])
        doc.add_paragraph()
    
    # –ê–Ω–∞–ª–∏–∑ –¥–µ—Ñ–µ–∫—Ç–æ–≤
    if not defects_df.empty:
        doc.add_paragraph("–ê–Ω–∞–ª–∏–∑ –¥–µ—Ñ–µ–∫—Ç–æ–≤", style='Heading 2')
        add_table_from_df(doc, defects_df)
        doc.add_paragraph()
    
    # –û–≥—Ä–∞–Ω–∏—á–µ–Ω–∏—è —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è
    if data["limitations"].strip():
        doc.add_paragraph("–û–≥—Ä–∞–Ω–∏—á–µ–Ω–∏—è —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è", style='Heading 2')
        # –ü—Ä–µ–æ–±—Ä–∞–∑—É–µ–º –º–∞—Ä–∫–∏—Ä–æ–≤–∞–Ω–Ω—ã–π —Å–ø–∏—Å–æ–∫ –≤ –Ω—É–º–µ—Ä–æ–≤–∞–Ω–Ω—ã–π
        lines = [line.strip() for line in data["limitations"].split('\n') if line.strip()]
        for line in lines:
            # –£–±–∏—Ä–∞–µ–º –º–∞—Ä–∫–µ—Ä—ã "-", "*" –µ—Å–ª–∏ –µ—Å—Ç—å
            clean_line = line.lstrip('-*‚Ä¢ ').strip()
            doc.add_paragraph(clean_line, style='List Number')
        doc.add_paragraph()
    
    # –í—ã–≤–æ–¥ –∏ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏
    doc.add_paragraph("–í—ã–≤–æ–¥ –∏ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏", style='Heading 2')
    
    if data["consequences"].strip():
        doc.add_paragraph("–ü–æ—Å–ª–µ–¥—Å—Ç–≤–∏—è –¥–µ—Ñ–µ–∫—Ç–æ–≤", style='Heading 3')
        doc.add_paragraph(data["consequences"])
    
    if data["conclusion"].strip():
        doc.add_paragraph("–í—ã–≤–æ–¥", style='Heading 3')
        doc.add_paragraph(data["conclusion"])
    
    if data["recommendations_detailed"].strip():
        doc.add_paragraph("–†–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏", style='Heading 3')
        doc.add_paragraph(data["recommendations_detailed"])
    doc.add_paragraph()
    
    # –ü–æ–¥–ø–∏—Å—å
    doc.add_paragraph("–ü–æ–¥–ø–∏—Å—å", style='Heading 2')
    
    signature_table = doc.add_table(rows=3, cols=2)
    signature_table.style = 'Table Grid'
    
    signature_fields = [
        ("–†–æ–ª—å", data["role"]),
        ("–§–ò–û", data["fullname"]),
        ("–î–∞—Ç–∞", data["signature_date"]),
    ]
    
    for i, (label, value) in enumerate(signature_fields):
        signature_table.cell(i, 0).text = label
        signature_table.cell(i, 1).text = str(value)
        for paragraph in signature_table.cell(i, 0).paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Calibri Light'
                run.font.size = Pt(10)
        for paragraph in signature_table.cell(i, 1).paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Calibri Light'
                run.font.size = Pt(10)
    
    set_col_width(signature_table, 0, 1800)
    set_col_width(signature_table, 1, 5400)
    
    # –°–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ –≤ –±—É—Ñ–µ—Ä
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# === –ì–ï–ù–ï–†–ê–¶–ò–Ø HTML –û–¢–ß–Å–¢–ê ===
def escape_html(text):
    if pd.isna(text) or text is None:
        return ""
    return (str(text)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#x27;"))

def generate_html_report(data, module_data_list, defects_df):
    # –í–∞–ª–∏–¥–∞—Ü–∏—è –º–µ—Ç—Ä–∏–∫
    validation_errors = []
    total_tc = data["total_tc"]
    pass_tc = data["pass"]
    fail_tc = data["fail"]
    
    if pass_tc + fail_tc != total_tc:
        validation_errors.append("‚ö†Ô∏è –°—É–º–º–∞ —Å—Ç–∞—Ç—É—Å–æ–≤ (PASS + FAIL) –Ω–µ —Ä–∞–≤–Ω–∞ –æ–±—â–µ–º—É –∫–æ–ª–∏—á–µ—Å—Ç–≤—É —Ç–µ—Å—Ç-–∫–µ–π—Å–æ–≤")
    
    # –ü–æ–¥–≥–æ—Ç–æ–≤–∫–∞ –¥–∞–Ω–Ω—ã—Ö –¥–ª—è –¥–∏–∞–≥—Ä–∞–º–º
    labels_results = ['PASS', 'FAIL']
    sizes_results = [pass_tc, fail_tc]
    colors_results = ['#4CAF50', '#F44336']
    
    labels_severity = []
    sizes_severity = []
    colors_severity_map = {'S1': '#F44336', 'S2': '#FF9800', 'S3': '#FFC107', 'S4': '#4CAF50'}
    
    if data['s1'] > 0:
        labels_severity.append('S1')
        sizes_severity.append(data['s1'])
    if data['s2'] > 0:
        labels_severity.append('S2')
        sizes_severity.append(data['s2'])
    
    # –°–æ–∑–¥–∞–Ω–∏–µ –¥–∏–∞–≥—Ä–∞–º–º –≤ base64
    def plot_to_base64(fig):
        buf = io.BytesIO()
        fig.savefig(buf, format='png', bbox_inches='tight', dpi=150)
        buf.seek(0)
        img_base64 = base64.b64encode(buf.read()).decode('utf-8')
        buf.close()
        plt.close(fig)
        return img_base64
    
    # –î–∏–∞–≥—Ä–∞–º–º–∞ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤
    fig1, ax1 = plt.subplots(figsize=(6, 4))
    ax1.pie(sizes_results, labels=labels_results, colors=colors_results, autopct='%1.1f%%', startangle=90)
    ax1.axis('equal')
    chart1_base64 = plot_to_base64(fig1)
    
    # –î–∏–∞–≥—Ä–∞–º–º–∞ —Å–µ—Ä—å—ë–∑–Ω–æ—Å—Ç–∏ (–µ—Å–ª–∏ –µ—Å—Ç—å –¥–µ—Ñ–µ–∫—Ç—ã)
    chart2_base64 = None
    if sizes_severity:
        fig2, ax2 = plt.subplots(figsize=(6, 4))
        colors_sev = [colors_severity_map.get(lbl, '#9E9E9E') for lbl in labels_severity]
        ax2.pie(sizes_severity, labels=labels_severity, colors=colors_sev, autopct='%1.1f%%', startangle=90)
        ax2.axis('equal')
        chart2_base64 = plot_to_base64(fig2)
    
    # –§–æ—Ä–º–∏—Ä–æ–≤–∞–Ω–∏–µ HTML
    html_content = f"""
    <!DOCTYPE html>
    <html lang="ru">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{escape_html(data['report_title'])}</title>
        <style>
            body {{
                font-family: 'Calibri', 'Segoe UI', Arial, sans-serif;
                line-height: 1.6;
                color: #333;
                max-width: 1200px;
                margin: 0 auto;
                padding: 20px;
                background-color: #f9f9f9;
            }}
            .container {{
                background: white;
                padding: 30px;
                border-radius: 8px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            }}
            h1 {{
                text-align: center;
                color: #1a365d;
                font-size: 24px;
                margin-bottom: 30px;
                border-bottom: 2px solid #4472C4;
                padding-bottom: 10px;
            }}
            h2 {{
                color: #4472C4;
                border-left: 4px solid #4472C4;
                padding-left: 10px;
                margin-top: 25px;
            }}
            h3 {{
                color: #5b616b;
                margin-top: 20px;
            }}
            .info-table {{
                width: 100%;
                border-collapse: collapse;
                margin: 15px 0;
                font-size: 14px;
            }}
            .info-table th, .info-table td {{
                border: 1px solid #ddd;
                padding: 8px 12px;
                text-align: left;
                vertical-align: top;
            }}
            .info-table th {{
                background-color: #f2f2f2;
                width: 25%;
                font-weight: bold;
            }}
            .status-pass {{
                background-color: #e8f5e9;
                color: #2e7d32;
                font-weight: bold;
            }}
            .status-fail {{
                background-color: #ffebee;
                color: #c62828;
                font-weight: bold;
            }}
            .severity-s1 {{
                background-color: #ffebee;
                color: #c62828;
                font-weight: bold;
            }}
            .severity-s2 {{
                background-color: #fff3e0;
                color: #e65100;
                font-weight: bold;
            }}
            .chart-container {{
                text-align: center;
                margin: 25px 0;
            }}
            .chart-container img {{
                max-width: 100%;
                height: auto;
                border: 1px solid #ddd;
                border-radius: 4px;
            }}
            .chart-caption {{
                font-size: 13px;
                color: #666;
                margin-top: 5px;
            }}
            .validation-error {{
                background-color: #ffebee;
                color: #c62828;
                padding: 10px;
                border-radius: 4px;
                margin: 15px 0;
                border-left: 4px solid #c62828;
            }}
            .limitations li {{
                margin-bottom: 5px;
            }}
            @media print {{
                body {{
                    background-color: white;
                    padding: 0;
                }}
                .container {{
                    box-shadow: none;
                    padding: 15px;
                }}
                .no-print {{
                    display: none;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <h1>{escape_html(data['report_title'])}</h1>
            
            <!-- –í–∞–ª–∏–¥–∞—Ü–∏—è -->
            {"".join([f'<div class="validation-error">{err}</div>' for err in validation_errors]) if validation_errors else ""}
            
            <!-- –û—Å–Ω–æ–≤–Ω–∞—è –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è -->
            <h2>1. –û—Å–Ω–æ–≤–Ω–∞—è –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è</h2>
            <table class="info-table">
                <tr><th>–ü—Ä–æ–µ–∫—Ç</th><td>{escape_html(data['project'])}</td></tr>
                <tr><th>–¢–∏–ø –ø—Ä–∏–ª–æ–∂–µ–Ω–∏—è</th><td>{escape_html(data['app_type'])}</td></tr>
                <tr><th>–í–µ—Ä—Å–∏—è</th><td>{escape_html(data['version'])}</td></tr>
                <tr><th>–ü–µ—Ä–∏–æ–¥ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è</th><td>{escape_html(data['test_period'])}</td></tr>
                <tr><th>–î–∞—Ç–∞ —Ñ–æ—Ä–º–∏—Ä–æ–≤–∞–Ω–∏—è –æ—Ç—á—ë—Ç–∞</th><td>{escape_html(data['report_date'])}</td></tr>
                <tr><th>–ò–Ω–∂–µ–Ω–µ—Ä –ø–æ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—é</th><td>{escape_html(data['engineer'])}</td></tr>
                <tr><th>–°—Ç–∞—Ç—É—Å —Ä–µ–ª–∏–∑–∞</th><td>{escape_html(data['release_status'])}</td></tr>
            </table>
            
            <!-- –ö—Ä–∞—Ç–∫–æ–µ —Ä–µ–∑—é–º–µ -->
            <h2>2. –ö—Ä–∞—Ç–∫–æ–µ —Ä–µ–∑—é–º–µ</h2>
            <table class="info-table">
                <tr><th>–ö–æ–ª–∏—á–µ—Å—Ç–≤–æ –¥–µ—Ñ–µ–∫—Ç–æ–≤ (S1)</th><td>{data['s1']}</td></tr>
                <tr><th>–ö–æ–ª–∏—á–µ—Å—Ç–≤–æ –¥–µ—Ñ–µ–∫—Ç–æ–≤ (S2)</th><td>{data['s2']}</td></tr>
                <tr><th>–í—Å–µ–≥–æ —Ç–µ—Å—Ç-–∫–µ–π—Å–æ–≤</th><td>{data['total_tc']}</td></tr>
                <tr><th>–ü—Ä–æ–π–¥–µ–Ω–æ —É—Å–ø–µ—à–Ω–æ</th><td>{data['pass']}</td></tr>
            </table>
            
            <h3>–†–∏—Å–∫–∏</h3>
            <p>{escape_html(data['risk'])}</p>
            
            <h3>–†–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏</h3>
            <p>{escape_html(data['recommendation'])}</p>
            
            <!-- –î–∏–∞–≥—Ä–∞–º–º—ã -->
            <h2>3. –î–∏–∞–≥—Ä–∞–º–º—ã</h2>
            <div class="chart-container">
                <img src="data:image/png;base64,{chart1_base64}" alt="–†–∞—Å–ø—Ä–µ–¥–µ–ª–µ–Ω–∏–µ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤">
                <div class="chart-caption">–†–∏—Å. 1. –†–∞—Å–ø—Ä–µ–¥–µ–ª–µ–Ω–∏–µ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è</div>
            </div>
            """
    
    if chart2_base64:
        html_content += f"""
            <div class="chart-container">
                <img src="data:image/png;base64,{chart2_base64}" alt="–†–∞—Å–ø—Ä–µ–¥–µ–ª–µ–Ω–∏–µ –¥–µ—Ñ–µ–∫—Ç–æ–≤ –ø–æ —Å–µ—Ä—å—ë–∑–Ω–æ—Å—Ç–∏">
                <div class="chart-caption">–†–∏—Å. 2. –†–∞—Å–ø—Ä–µ–¥–µ–ª–µ–Ω–∏–µ –¥–µ—Ñ–µ–∫—Ç–æ–≤ –ø–æ —Å–µ—Ä—å—ë–∑–Ω–æ—Å—Ç–∏</div>
            </div>
            """
    
    # –ö–æ–Ω—Ç–µ–∫—Å—Ç —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è
    html_content += f"""
            <h2>4. –ö–æ–Ω—Ç–µ–∫—Å—Ç —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è</h2>
            <table class="info-table">
                <tr><th>–£—Å—Ç—Ä–æ–π—Å—Ç–≤–æ / –ë—Ä–∞—É–∑–µ—Ä</th><td>{escape_html(data['device_browser'])}</td></tr>
                <tr><th>–û–° / –ü–ª–∞—Ç—Ñ–æ—Ä–º–∞</th><td>{escape_html(data['os_platform'])}</td></tr>
                <tr><th>–°–±–æ—Ä–∫–∞</th><td>{escape_html(data['build'])}</td></tr>
                <tr><th>URL –æ–∫—Ä—É–∂–µ–Ω–∏—è</th><td>{escape_html(data['env_url'])}</td></tr>
                <tr><th>–ò–Ω—Å—Ç—Ä—É–º–µ–Ω—Ç—ã</th><td>{escape_html(data['tools'])}</td></tr>
            </table>
            
            <h3>–ú–µ—Ç–æ–¥–æ–ª–æ–≥–∏—è —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è</h3>
            <p>{escape_html(data['methodology'])}</p>
            """
    
    # –†–µ–∑—É–ª—å—Ç–∞—Ç—ã –ø–æ –º–æ–¥—É–ª—è–º
    html_content += f"""
            <h2>5. –†–µ–∑—É–ª—å—Ç–∞—Ç—ã —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è –ø–æ –º–æ–¥—É–ª—è–º</h2>
            """
    
    for module in module_data_list:
        html_content += f"""
            <h3>{escape_html(module['title'])}</h3>
            <table class="info-table">
                <thead>
                    <tr>
                        <th>ID</th>
                        <th>–°—Ü–µ–Ω–∞—Ä–∏–π</th>
                        <th>–°—Ç–∞—Ç—É—Å</th>
                        <th>–ö–æ–º–º–µ–Ω—Ç–∞—Ä–∏–π</th>
                    </tr>
                </thead>
                <tbody>
            """
        for _, row in module['df'].iterrows():
            status_class = "status-pass" if str(row['–°—Ç–∞—Ç—É—Å']).strip().upper() == "PASS" else "status-fail"
            html_content += f"""
                    <tr>
                        <td>{escape_html(row['ID'])}</td>
                        <td>{escape_html(row['–°—Ü–µ–Ω–∞—Ä–∏–π'])}</td>
                        <td class="{status_class}">{escape_html(row['–°—Ç–∞—Ç—É—Å'])}</td>
                        <td>{escape_html(row['–ö–æ–º–º–µ–Ω—Ç–∞—Ä–∏–π'])}</td>
                    </tr>
            """
        html_content += """
                </tbody>
            </table>
            """
    
    # –ê–Ω–∞–ª–∏–∑ –¥–µ—Ñ–µ–∫—Ç–æ–≤
    if not defects_df.empty:
        html_content += f"""
            <h2>6. –ê–Ω–∞–ª–∏–∑ –¥–µ—Ñ–µ–∫—Ç–æ–≤</h2>
            <table class="info-table">
                <thead>
                    <tr>
                        <th>ID</th>
                        <th>–ú–æ–¥—É–ª—å</th>
                        <th>–ó–∞–≥–æ–ª–æ–≤–æ–∫</th>
                        <th>–°–µ—Ä—å—ë–∑–Ω–æ—Å—Ç—å</th>
                        <th>–°—Ç–∞—Ç—É—Å</th>
                    </tr>
                </thead>
                <tbody>
            """
        for _, row in defects_df.iterrows():
            sev_class = f"severity-{str(row['–°–µ—Ä—å—ë–∑–Ω–æ—Å—Ç—å']).lower()}" if pd.notna(row['–°–µ—Ä—å—ë–∑–Ω–æ—Å—Ç—å']) else ""
            html_content += f"""
                    <tr>
                        <td>{escape_html(row['ID'])}</td>
                        <td>{escape_html(row['–ú–æ–¥—É–ª—å'])}</td>
                        <td>{escape_html(row['–ó–∞–≥–æ–ª–æ–≤–æ–∫'])}</td>
                        <td class="{sev_class}">{escape_html(row['–°–µ—Ä—å—ë–∑–Ω–æ—Å—Ç—å'])}</td>
                        <td>{escape_html(row['–°—Ç–∞—Ç—É—Å'])}</td>
                    </tr>
            """
        html_content += """
                </tbody>
            </table>
            """
    
    # –û–≥—Ä–∞–Ω–∏—á–µ–Ω–∏—è —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è
    if data["limitations"].strip():
        html_content += f"""
            <h2>7. –û–≥—Ä–∞–Ω–∏—á–µ–Ω–∏—è —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è</h2>
            <div class="limitations">
                <ol>
        """
        lines = [line.strip() for line in data["limitations"].split('\n') if line.strip()]
        for line in lines:
            clean_line = line.lstrip('-*‚Ä¢ ').strip()
            html_content += f"<li>{escape_html(clean_line)}</li>"
        html_content += """
                </ol>
            </div>
            """
    
    # –í—ã–≤–æ–¥ –∏ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏
    html_content += f"""
            <h2>8. –í—ã–≤–æ–¥ –∏ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏</h2>
            """
    
    if data["consequences"].strip():
        html_content += f"""
            <h3>–ü–æ—Å–ª–µ–¥—Å—Ç–≤–∏—è –¥–µ—Ñ–µ–∫—Ç–æ–≤</h3>
            <p>{escape_html(data['consequences'])}</p>
            """
    
    if data["conclusion"].strip():
        html_content += f"""
            <h3>–í—ã–≤–æ–¥</h3>
            <p>{escape_html(data['conclusion'])}</p>
            """
    
    if data["recommendations_detailed"].strip():
        html_content += f"""
            <h3>–†–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏</h3>
            <p>{escape_html(data['recommendations_detailed'])}</p>
            """
    
    # –ü–æ–¥–ø–∏—Å—å
    html_content += f"""
            <h2>9. –ü–æ–¥–ø–∏—Å—å</h2>
            <table class="info-table">
                <tr><th>–†–æ–ª—å</th><td>{escape_html(data['role'])}</td></tr>
                <tr><th>–§–ò–û</th><td>{escape_html(data['fullname'])}</td></tr>
                <tr><th>–î–∞—Ç–∞</th><td>{escape_html(data['signature_date'])}</td></tr>
            </table>
            
            <div class="no-print" style="margin-top: 30px; text-align: center; color: #666; font-size: 12px;">
                –û—Ç—á—ë—Ç —Å–≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞–Ω –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏ —á–µ—Ä–µ–∑ –∏–Ω—Å—Ç—Ä—É–º–µ–Ω—Ç —Ç–µ—Å—Ç–æ–≤–æ–π –æ—Ç—á—ë—Ç–Ω–æ—Å—Ç–∏
            </div>
        </div>
    </body>
    </html>
    """
    
    return html_content

# === –ì–ï–ù–ï–†–ê–¶–ò–Ø XLSX –û–¢–ß–Å–¢–ê ===
def generate_xlsx_single_sheet(data, module_data_list, defects_df):
    wb = Workbook()
    ws = wb.active
    ws.title = "–û—Ç—á—ë—Ç –æ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏–∏"
    
    # –°—Ç–∏–ª–∏
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(name='Calibri Light', size=11, bold=True, color="FFFFFF")
    normal_font = Font(name='Calibri Light', size=11)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    align_left = Alignment(horizontal='left', vertical='top', wrap_text=True)
    align_center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    
    # –ó–∞–≥–æ–ª–æ–≤–æ–∫
    ws.merge_cells('A1:D1')
    ws['A1'] = data["report_title"]
    ws['A1'].font = Font(name='Calibri Light', size=16, bold=True)
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 30
    ws.column_dimensions['A'].width = 15
    ws.column_dimensions['B'].width = 40
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 40
    
    row = 3
    
    # –û—Å–Ω–æ–≤–Ω–∞—è –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è
    ws.cell(row=row, column=1, value="–û–°–ù–û–í–ù–ê–Ø –ò–ù–§–û–†–ú–ê–¶–ò–Ø")
    ws.cell(row=row, column=1).font = Font(name='Calibri Light', size=14, bold=True, color="4472C4")
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
    row += 1
    
    info_fields = [
        ("–ü—Ä–æ–µ–∫—Ç", data["project"]),
        ("–¢–∏–ø –ø—Ä–∏–ª–æ–∂–µ–Ω–∏—è", data["app_type"]),
        ("–í–µ—Ä—Å–∏—è", data["version"]),
        ("–ü–µ—Ä–∏–æ–¥ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è", data["test_period"]),
        ("–î–∞—Ç–∞ —Ñ–æ—Ä–º–∏—Ä–æ–≤–∞–Ω–∏—è –æ—Ç—á—ë—Ç–∞", data["report_date"]),
        ("–ò–Ω–∂–µ–Ω–µ—Ä –ø–æ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—é", data["engineer"]),
        ("–°—Ç–∞—Ç—É—Å —Ä–µ–ª–∏–∑–∞", data["release_status"]),
    ]
    
    for label, value in info_fields:
        ws.cell(row=row, column=1, value=label).font = Font(name='Calibri Light', size=11, bold=True)
        ws.cell(row=row, column=2, value=value).font = normal_font
        ws.cell(row=row, column=1).border = border
        ws.cell(row=row, column=2).border = border
        ws.merge_cells(start_row=row, start_column=3, end_row=row, end_column=4)
        row += 1
    
    row += 1
    
    # –ö—Ä–∞—Ç–∫–æ–µ —Ä–µ–∑—é–º–µ
    ws.cell(row=row, column=1, value="–ö–†–ê–¢–ö–û–ï –†–ï–ó–Æ–ú–ï")
    ws.cell(row=row, column=1).font = Font(name='Calibri Light', size=14, bold=True, color="4472C4")
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
    row += 1
    
    metrics = [
        ("–ö–æ–ª–∏—á–µ—Å—Ç–≤–æ –¥–µ—Ñ–µ–∫—Ç–æ–≤ (S1)", data["s1"]),
        ("–ö–æ–ª–∏—á–µ—Å—Ç–≤–æ –¥–µ—Ñ–µ–∫—Ç–æ–≤ (S2)", data["s2"]),
        ("–í—Å–µ–≥–æ —Ç–µ—Å—Ç-–∫–µ–π—Å–æ–≤", data["total_tc"]),
        ("–ü—Ä–æ–π–¥–µ–Ω–æ —É—Å–ø–µ—à–Ω–æ", data["pass"]),
        ("–£–ø–∞–ª–æ", data["fail"]),
    ]
    
    for label, value in metrics:
        ws.cell(row=row, column=1, value=label).font = Font(name='Calibri Light', size=11, bold=True)
        ws.cell(row=row, column=2, value=value).font = normal_font
        ws.cell(row=row, column=1).border = border
        ws.cell(row=row, column=2).border = border
        ws.merge_cells(start_row=row, start_column=3, end_row=row, end_column=4)
        row += 1
    
    row += 1
    
    # –†–∏—Å–∫–∏ –∏ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏
    if data["risk"].strip():
        ws.cell(row=row, column=1, value="–†–ò–°–ö–ò").font = Font(name='Calibri Light', size=12, bold=True)
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        row += 1
        ws.cell(row=row, column=1, value=data["risk"]).font = normal_font
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        ws.cell(row=row, column=1).alignment = align_left
        row += 2
    
    if data["recommendation"].strip():
        ws.cell(row=row, column=1, value="–†–ï–ö–û–ú–ï–ù–î–ê–¶–ò–ò").font = Font(name='Calibri Light', size=12, bold=True)
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        row += 1
        ws.cell(row=row, column=1, value=data["recommendation"]).font = normal_font
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        ws.cell(row=row, column=1).alignment = align_left
        row += 2
    
    # –†–µ–∑—É–ª—å—Ç–∞—Ç—ã –ø–æ –º–æ–¥—É–ª—è–º
    ws.cell(row=row, column=1, value="–†–ï–ó–£–õ–¨–¢–ê–¢–´ –ü–û –ú–û–î–£–õ–Ø–ú")
    ws.cell(row=row, column=1).font = Font(name='Calibri Light', size=14, bold=True, color="4472C4")
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
    row += 1
    
    for module in module_data_list:
        ws.cell(row=row, column=1, value=module["title"]).font = Font(name='Calibri Light', size=12, bold=True)
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        row += 1
        
        # –ó–∞–≥–æ–ª–æ–≤–∫–∏ —Ç–∞–±–ª–∏—Ü—ã –º–æ–¥—É–ª—è
        headers = ["ID", "–°—Ü–µ–Ω–∞—Ä–∏–π", "–°—Ç–∞—Ç—É—Å", "–ö–æ–º–º–µ–Ω—Ç–∞—Ä–∏–π"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border
            cell.alignment = align_center
        row += 1
        
        # –î–∞–Ω–Ω—ã–µ –º–æ–¥—É–ª—è
        for _, r in module["df"].iterrows():
            ws.cell(row=row, column=1, value=r["ID"]).font = normal_font
            ws.cell(row=row, column=2, value=r["–°—Ü–µ–Ω–∞—Ä–∏–π"]).font = normal_font
            status_cell = ws.cell(row=row, column=3, value=r["–°—Ç–∞—Ç—É—Å"])
            status_cell.font = normal_font
            
            # –¶–≤–µ—Ç —Å—Ç–∞—Ç—É—Å–∞
            status_val = str(r["–°—Ç–∞—Ç—É—Å"]).strip().upper() if pd.notna(r["–°—Ç–∞—Ç—É—Å"]) else ""
            if status_val == "PASS":
                status_cell.fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
                status_cell.font = Font(name='Calibri Light', size=11, bold=True, color="006100")
            elif status_val == "FAIL":
                status_cell.fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
                status_cell.font = Font(name='Calibri Light', size=11, bold=True, color="9C0006")
            
            ws.cell(row=row, column=4, value=r["–ö–æ–º–º–µ–Ω—Ç–∞—Ä–∏–π"] if pd.notna(r["–ö–æ–º–º–µ–Ω—Ç–∞—Ä–∏–π"]) else "").font = normal_font
            
            for col in range(1, 5):
                ws.cell(row=row, column=col).border = border
                ws.cell(row=row, column=col).alignment = align_left
            
            row += 1
        
        row += 1
    
    # –ê–Ω–∞–ª–∏–∑ –¥–µ—Ñ–µ–∫—Ç–æ–≤
    if not defects_df.empty:
        ws.cell(row=row, column=1, value="–ê–ù–ê–õ–ò–ó –î–ï–§–ï–ö–¢–û–í")
        ws.cell(row=row, column=1).font = Font(name='Calibri Light', size=14, bold=True, color="4472C4")
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        row += 1
        
        # –ó–∞–≥–æ–ª–æ–≤–∫–∏ –¥–µ—Ñ–µ–∫—Ç–æ–≤
        defect_headers = ["ID", "–ú–æ–¥—É–ª—å", "–ó–∞–≥–æ–ª–æ–≤–æ–∫", "–°–µ—Ä—å—ë–∑–Ω–æ—Å—Ç—å"]
        for col, header in enumerate(defect_headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border
            cell.alignment = align_center
        row += 1
        
        # –î–∞–Ω–Ω—ã–µ –¥–µ—Ñ–µ–∫—Ç–æ–≤
        for _, r in defects_df.iterrows():
            ws.cell(row=row, column=1, value=r["ID"]).font = normal_font
            ws.cell(row=row, column=2, value=r["–ú–æ–¥—É–ª—å"]).font = normal_font
            ws.cell(row=row, column=3, value=r["–ó–∞–≥–æ–ª–æ–≤–æ–∫"]).font = normal_font
            
            sev_cell = ws.cell(row=row, column=4, value=r["–°–µ—Ä—å—ë–∑–Ω–æ—Å—Ç—å"])
            sev_cell.font = normal_font
            
            # –¶–≤–µ—Ç —Å–µ—Ä—å—ë–∑–Ω–æ—Å—Ç–∏
            sev_val = str(r["–°–µ—Ä—å—ë–∑–Ω–æ—Å—Ç—å"]).strip().upper() if pd.notna(r["–°–µ—Ä—å—ë–∑–Ω–æ—Å—Ç—å"]) else ""
            if sev_val == "S1":
                sev_cell.fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
                sev_cell.font = Font(name='Calibri Light', size=11, bold=True, color="9C0006")
            elif sev_val == "S2":
                sev_cell.fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
                sev_cell.font = Font(name='Calibri Light', size=11, bold=True, color="9C5700")
            
            for col in range(1, 5):
                ws.cell(row=row, column=col).border = border
                ws.cell(row=row, column=col).alignment = align_left
            
            row += 1
        
        row += 1
    
    # –û–≥—Ä–∞–Ω–∏—á–µ–Ω–∏—è
    if data["limitations"].strip():
        ws.cell(row=row, column=1, value="–û–ì–†–ê–ù–ò–ß–ï–ù–ò–Ø –¢–ï–°–¢–ò–†–û–í–ê–ù–ò–Ø")
        ws.cell(row=row, column=1).font = Font(name='Calibri Light', size=14, bold=True, color="4472C4")
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        row += 1
        
        lines = [line.strip() for line in data["limitations"].split('\n') if line.strip()]
        for i, line in enumerate(lines, 1):
            clean_line = line.lstrip('-*‚Ä¢ ').strip()
            ws.cell(row=row, column=1, value=f"{i}. {clean_line}").font = normal_font
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
            ws.cell(row=row, column=1).alignment = align_left
            row += 1
        
        row += 1
    
    # –í—ã–≤–æ–¥ –∏ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏
    ws.cell(row=row, column=1, value="–í–´–í–û–î –ò –†–ï–ö–û–ú–ï–ù–î–ê–¶–ò–ò")
    ws.cell(row=row, column=1).font = Font(name='Calibri Light', size=14, bold=True, color="4472C4")
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
    row += 1
    
    if data["consequences"].strip():
        ws.cell(row=row, column=1, value="–ü–æ—Å–ª–µ–¥—Å—Ç–≤–∏—è –¥–µ—Ñ–µ–∫—Ç–æ–≤").font = Font(name='Calibri Light', size=12, bold=True)
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        row += 1
        ws.cell(row=row, column=1, value=data["consequences"]).font = normal_font
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        ws.cell(row=row, column=1).alignment = align_left
        row += 2
    
    if data["conclusion"].strip():
        ws.cell(row=row, column=1, value="–í—ã–≤–æ–¥").font = Font(name='Calibri Light', size=12, bold=True)
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        row += 1
        ws.cell(row=row, column=1, value=data["conclusion"]).font = normal_font
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        ws.cell(row=row, column=1).alignment = align_left
        row += 2
    
    if data["recommendations_detailed"].strip():
        ws.cell(row=row, column=1, value="–†–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏").font = Font(name='Calibri Light', size=12, bold=True)
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        row += 1
        ws.cell(row=row, column=1, value=data["recommendations_detailed"]).font = normal_font
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        ws.cell(row=row, column=1).alignment = align_left
        row += 2
    
    # –ü–æ–¥–ø–∏—Å—å
    ws.cell(row=row, column=1, value="–ü–û–î–ü–ò–°–¨")
    ws.cell(row=row, column=1).font = Font(name='Calibri Light', size=14, bold=True, color="4472C4")
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
    row += 1
    
    signature_fields = [
        ("–†–æ–ª—å", data["role"]),
        ("–§–ò–û", data["fullname"]),
        ("–î–∞—Ç–∞", data["signature_date"]),
    ]
    
    for label, value in signature_fields:
        ws.cell(row=row, column=1, value=label).font = Font(name='Calibri Light', size=11, bold=True)
        ws.cell(row=row, column=2, value=value).font = normal_font
        ws.cell(row=row, column=1).border = border
        ws.cell(row=row, column=2).border = border
        ws.merge_cells(start_row=row, start_column=3, end_row=row, end_column=4)
        row += 1
    
    # –ê–≤—Ç–æ–ø–æ–¥–±–æ—Ä –≤—ã—Å–æ—Ç—ã —Å—Ç—Ä–æ–∫
    for row_idx in range(1, row):
        ws.row_dimensions[row_idx].height = 15
    
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer

# === –î–ê–ù–ù–´–ï –ü–û –£–ú–û–õ–ß–ê–ù–ò–Æ ===
default_modules = [
    {
        "title": "–ê–≤—Ç–æ—Ä–∏–∑–∞—Ü–∏—è –∏ —Ä–µ–≥–∏—Å—Ç—Ä–∞—Ü–∏—è",
        "df": pd.DataFrame([
            {"ID": "AUTH-001", "–°—Ü–µ–Ω–∞—Ä–∏–π": "–í—Ö–æ–¥ –ø–æ –≤–∞–ª–∏–¥–Ω—ã–º –¥–∞–Ω–Ω—ã–º", "–°—Ç–∞—Ç—É—Å": "PASS", "–ö–æ–º–º–µ–Ω—Ç–∞—Ä–∏–π": ""},
            {"ID": "AUTH-002", "–°—Ü–µ–Ω–∞—Ä–∏–π": "–í—Ö–æ–¥ –ø–æ –Ω–µ–≤–∞–ª–∏–¥–Ω—ã–º –¥–∞–Ω–Ω—ã–º", "–°—Ç–∞—Ç—É—Å": "PASS", "–ö–æ–º–º–µ–Ω—Ç–∞—Ä–∏–π": ""},
            {"ID": "AUTH-003", "–°—Ü–µ–Ω–∞—Ä–∏–π": "–í–æ—Å—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω–∏–µ –ø–∞—Ä–æ–ª—è", "–°—Ç–∞—Ç—É—Å": "FAIL", "–ö–æ–º–º–µ–Ω—Ç–∞—Ä–∏–π": "–ü–∏—Å—å–º–æ –Ω–µ –ø—Ä–∏—Ö–æ–¥–∏—Ç –Ω–∞ –ø–æ—á—Ç—É"},
        ])
    },
    {
        "title": "–ü–æ–∏—Å–∫ —Ç–æ–≤–∞—Ä–æ–≤",
        "df": pd.DataFrame([
            {"ID": "SEARCH-001", "–°—Ü–µ–Ω–∞—Ä–∏–π": "–ü–æ–∏—Å–∫ –ø–æ —Ç–æ—á–Ω–æ–º—É –Ω–∞–∑–≤–∞–Ω–∏—é", "–°—Ç–∞—Ç—É—Å": "PASS", "–ö–æ–º–º–µ–Ω—Ç–∞—Ä–∏–π": ""},
            {"ID": "SEARCH-002", "–°—Ü–µ–Ω–∞—Ä–∏–π": "–ü–æ–∏—Å–∫ —Å –æ–ø–µ—á–∞—Ç–∫–æ–π", "–°—Ç–∞—Ç—É—Å": "FAIL", "–ö–æ–º–º–µ–Ω—Ç–∞—Ä–∏–π": "Fuzzy search –Ω–µ —Ä–∞–±–æ—Ç–∞–µ—Ç"},
        ])
    }
]

default_defects = pd.DataFrame([
    {"ID": "BUG-SEC-001", "–ú–æ–¥—É–ª—å": "–ê–≤—Ç–æ—Ä–∏–∑–∞—Ü–∏—è", "–ó–∞–≥–æ–ª–æ–≤–æ–∫": "SQL-–∏–Ω—ä–µ–∫—Ü–∏—è –≤ –ø–æ–ª–µ –ª–æ–≥–∏–Ω–∞", "–°–µ—Ä—å—ë–∑–Ω–æ—Å—Ç—å": "S1", "–°—Ç–∞—Ç—É—Å": "–û—Ç–∫—Ä—ã—Ç"},
    {"ID": "BUG-SEC-002", "–ú–æ–¥—É–ª—å": "–ê–≤—Ç–æ—Ä–∏–∑–∞—Ü–∏—è", "–ó–∞–≥–æ–ª–æ–≤–æ–∫": "–ü–∞—Ä–æ–ª–∏ —Ö—Ä–∞–Ω—è—Ç—Å—è –≤ –æ—Ç–∫—Ä—ã—Ç–æ–º –≤–∏–¥–µ", "–°–µ—Ä—å—ë–∑–Ω–æ—Å—Ç—å": "S1", "–°—Ç–∞—Ç—É—Å": "–û—Ç–∫—Ä—ã—Ç"},
    {"ID": "BUG-SEARCH-001", "–ú–æ–¥—É–ª—å": "–ü–æ–∏—Å–∫", "–ó–∞–≥–æ–ª–æ–≤–æ–∫": "–ù–µ—Ç –ø–æ–¥–¥–µ—Ä–∂–∫–∏ –æ–ø–µ—á–∞—Ç–æ–∫", "–°–µ—Ä—å—ë–∑–Ω–æ—Å—Ç—å": "S2", "–°—Ç–∞—Ç—É—Å": "–û—Ç–∫—Ä—ã—Ç"},
])

# === –ò–ù–¢–ï–†–§–ï–ô–° STREAMLIT ===
st.set_page_config(page_title="–ì–µ–Ω–µ—Ä–∞—Ç–æ—Ä –æ—Ç—á—ë—Ç–∞", layout="wide")

# === –ó–ê–ì–†–£–ó–ö–ê –ß–ï–†–ù–û–í–ò–ö–ê (–í–ù–ï –§–û–†–ú–´!) ===
st.title("üìÑ –ì–µ–Ω–µ—Ä–∞—Ç–æ—Ä –ø—Ä–æ—Ñ–µ—Å—Å–∏–æ–Ω–∞–ª—å–Ω—ã—Ö —Ç–µ—Å—Ç–æ–≤—ã—Ö –æ—Ç—á—ë—Ç–æ–≤")
st.markdown("""
–°–æ–∑–¥–∞–≤–∞–π—Ç–µ –æ—Ç—á—ë—Ç—ã –æ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏–∏ –≤ —Ç—Ä—ë—Ö —Ñ–æ—Ä–º–∞—Ç–∞—Ö (DOCX, HTML, XLSX) —Å —Å–æ–±–ª—é–¥–µ–Ω–∏–µ–º –∫–æ—Ä–ø–æ—Ä–∞—Ç–∏–≤–Ω—ã—Ö —Å—Ç–∞–Ω–¥–∞—Ä—Ç–æ–≤.
–°–æ—Ö—Ä–∞–Ω—è–π—Ç–µ —á–µ—Ä–Ω–æ–≤–∏–∫–∏ –∏ –≤–æ–∑–≤—Ä–∞—â–∞–π—Ç–µ—Å—å –∫ –Ω–∏–º –ø–æ–∑–∂–µ!
""")

# –ö–Ω–æ–ø–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ —á–µ—Ä–Ω–æ–≤–∏–∫–∞ –í–ù–ï —Ñ–æ—Ä–º—ã
uploaded_file = st.file_uploader(
    "‚¨ÜÔ∏è –ó–∞–≥—Ä—É–∑–∏—Ç—å —á–µ—Ä–Ω–æ–≤–∏–∫ (.json)",
    type=["json"],
    label_visibility="visible",
    key="draft_uploader_outside_form"
)
if uploaded_file is not None:
    content = uploaded_file.read().decode("utf-8")
    restored_data, restored_modules, restored_defects, saved_at = load_draft(content)
    if restored_data is not None:
        st.session_state.draft_data = restored_data
        st.session_state.draft_modules = restored_modules
        st.session_state.draft_defects = restored_defects
        st.session_state.draft_saved_at = saved_at
        st.success(f"‚úÖ –ß–µ—Ä–Ω–æ–≤–∏–∫ –∑–∞–≥—Ä—É–∂–µ–Ω! –°–æ—Ö—Ä–∞–Ω—ë–Ω: {saved_at}")
        st.rerun()

# === –ü–û–î–ì–û–¢–û–í–ö–ê –î–ê–ù–ù–´–• –î–õ–Ø –§–û–†–ú–´ ===
if "draft_data" in st.session_state and st.session_state.draft_data is not None:
    draft_data = st.session_state.draft_data
    draft_modules = st.session_state.draft_modules
    draft_defects = st.session_state.draft_defects
    draft_saved_at = st.session_state.draft_saved_at
    
    # –ü–æ–¥—Å—Ç–∞–≤–ª—è–µ–º –∑–Ω–∞—á–µ–Ω–∏—è –∏–∑ —á–µ—Ä–Ω–æ–≤–∏–∫–∞
    report_title_val = draft_data.get("report_title", "–û—Ç—á—ë—Ç –æ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏–∏ –º–æ–±–∏–ª—å–Ω–æ–≥–æ –ø—Ä–∏–ª–æ–∂–µ–Ω–∏—è –õ–µ–º–∞–Ω–∞ –ü–†–û")
    project_val = draft_data.get("project", "–õ–µ–º–∞–Ω–∞ –ü–†–û")
    app_type_val = draft_data.get("app_type", "–ú–æ–±–∏–ª—å–Ω–æ–µ")
    version_val = draft_data.get("version", "241006.001")
    test_period_val = draft_data.get("test_period", "29‚Äì30 –Ω–æ—è–±—Ä—è 2025 –≥.")
    report_date_val = draft_data.get("report_date", "30 –Ω–æ—è–±—Ä—è 2025 –≥.")
    engineer_val = draft_data.get("engineer", "–ß–µ—Ä–∫–∞—Å–æ–≤ –ò–≥–æ—Ä—å")
    
    release_status_val = draft_data.get("release_status", "–ù–ï –†–ï–ö–û–ú–ï–ù–î–û–í–ê–ù –ö –í–´–ü–£–°–ö–£")
    s1_val = draft_data.get("s1", 2)
    s2_val = draft_data.get("s2", 1)
    total_tc_val = draft_data.get("total_tc", 72)
    pass_tc_val = draft_data.get("pass", 69)
    fail_tc_val = draft_data.get("fail", 3)
    risk_val = draft_data.get("risk", "–£—è–∑–≤–∏–º–æ—Å—Ç–∏ –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏ –ø–æ–∑–≤–æ–ª—è—é—Ç –Ω–∞—Ä—É—à–∏—Ç–µ–ª—é –ø–æ–ª—É—á–∏—Ç—å –¥–æ—Å—Ç—É–ø –∫ –¥–∞–Ω–Ω—ã–º –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª–µ–π –∏ –≤—ã–∑–≤–∞—Ç—å –æ—Ç–∫–∞–∑ –≤ –æ–±—Å–ª—É–∂–∏–≤–∞–Ω–∏–∏.")
    recommendation_val = draft_data.get("recommendation", "–†–µ–ª–∏–∑ –≤–æ–∑–º–æ–∂–µ–Ω —Ç–æ–ª—å–∫–æ –ø–æ—Å–ª–µ —É—Å—Ç—Ä–∞–Ω–µ–Ω–∏—è –≤—Å–µ—Ö S1/S2 –¥–µ—Ñ–µ–∫—Ç–æ–≤ –∏ –ø–æ–≤—Ç–æ—Ä–Ω–æ–≥–æ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è.")
    
    device_browser_val = draft_data.get("device_browser", "Xiaomi 12")
    os_platform_val = draft_data.get("os_platform", "Android 15")
    build_val = draft_data.get("build", "lemanna-pro_241006.001.apk")
    env_url_val = draft_data.get("env_url", "https://test.lemanna.pro")
    tools_val = draft_data.get("tools", "Postman (API), Burp Suite (–±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç—å), Jira (–±–∞–≥-—Ç—Ä–µ–∫–∏–Ω–≥)")
    methodology_val = draft_data.get("methodology", "–†—É—á–Ω–æ–µ —Ñ—É–Ω–∫—Ü–∏–æ–Ω–∞–ª—å–Ω–æ–µ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ + –ø—Ä–æ–≤–µ—Ä–∫–∞ –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏")
    
    consequences_val = draft_data.get("consequences", "- S1 –¥–µ—Ñ–µ–∫—Ç—ã –ø–æ–∑–≤–æ–ª—è—é—Ç –∑–ª–æ—É–º—ã—à–ª–µ–Ω–Ω–∏–∫—É –ø–æ–ª—É—á–∏—Ç—å –¥–∞–Ω–Ω—ã–µ –¥—Ä—É–≥–∏—Ö –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª–µ–π –∏–ª–∏ –≤—ã–≤–µ—Å—Ç–∏ –ø—Ä–∏–ª–æ–∂–µ–Ω–∏–µ –∏–∑ —Å—Ç—Ä–æ—è.\n- S2 –¥–µ—Ñ–µ–∫—Ç —Å–Ω–∏–∂–∞–µ—Ç —é–∑–∞–±–∏–ª–∏—Ç–∏: –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª–∏ –Ω–µ –Ω–∞–π–¥—É—Ç —Ç–æ–≤–∞—Ä –ø—Ä–∏ –æ–ø–µ—á–∞—Ç–∫–µ.")
    limitations_val = draft_data.get("limitations", "1. –ù–µ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–ª–∞—Å—å –æ–ø–ª–∞—Ç–∞ —á–µ—Ä–µ–∑ Apple Pay (—É—Å—Ç—Ä–æ–π—Å—Ç–≤–æ Android).\n2. –ù–µ –ø—Ä–æ–≤–µ—Ä–µ–Ω–∞ —Å–∏–Ω—Ö—Ä–æ–Ω–∏–∑–∞—Ü–∏—è —Å 1–° (–Ω–µ—Ç –¥–æ—Å—Ç—É–ø–∞ –∫ –∏–Ω—Ç–µ–≥—Ä–∞—Ü–∏–æ–Ω–Ω–æ–º—É —Å—Ç–µ–Ω–¥—É).\n3. –ù–µ –ø—Ä–æ–≤–µ–¥–µ–Ω–æ –Ω–∞–≥—Ä—É–∑–æ—á–Ω–æ–µ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ (–æ–≥—Ä–∞–Ω–∏—á–µ–Ω–∏–µ –ø–æ –≤—Ä–µ–º–µ–Ω–∏).")
    conclusion_val = draft_data.get("conclusion", "–°–±–æ—Ä–∫–∞ 241006.001 —Å–æ–¥–µ—Ä–∂–∏—Ç –∫—Ä–∏—Ç–∏—á–µ—Å–∫–∏–µ —É—è–∑–≤–∏–º–æ—Å—Ç–∏ –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏, –¥–µ–ª–∞—é—â–∏–µ –µ—ë –Ω–µ–ø—Ä–∏–≥–æ–¥–Ω–æ–π –¥–ª—è –≤—ã–ø—É—Å–∫–∞ –≤ production. –ù–∞–ª–∏—á–∏–µ S1 –¥–µ—Ñ–µ–∫—Ç–æ–≤ –Ω–∞—Ä—É—à–∞–µ—Ç –±–∞–∑–æ–≤—ã–µ –ø—Ä–∏–Ω—Ü–∏–ø—ã –∑–∞—â–∏—Ç—ã –¥–∞–Ω–Ω—ã—Ö –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª–µ–π.")
    recommendations_detailed_val = draft_data.get("recommendations_detailed", "–ù–µ–º–µ–¥–ª–µ–Ω–Ω–æ –∏—Å–ø—Ä–∞–≤–∏—Ç—å —É—è–∑–≤–∏–º–æ—Å—Ç–∏ BUG-SEC-001 –∏ BUG-SEC-002.\n–†–µ–∞–ª–∏–∑–æ–≤–∞—Ç—å fuzzy search –¥–ª—è –ø–æ–≤—ã—à–µ–Ω–∏—è —é–∑–∞–±–∏–ª–∏—Ç–∏ (BUG-SEARCH-001).\n–ü—Ä–æ–≤–µ—Å—Ç–∏ –ø–æ–≤—Ç–æ—Ä–Ω–æ–µ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ –ø–æ—Å–ª–µ —Ñ–∏–∫—Å–æ–≤ —Å —Ñ–æ–∫—É—Å–æ–º –Ω–∞:\n- –ü–æ–≤—Ç–æ—Ä–Ω—É—é –ø—Ä–æ–≤–µ—Ä–∫—É –ø–æ–ª–µ–π –≤–≤–æ–¥–∞ –Ω–∞ –∏–Ω—ä–µ–∫—Ü–∏–∏\n- –¢–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ —Å—Ü–µ–Ω–∞—Ä–∏–µ–≤ –ø–æ–∏—Å–∫–∞ —Å –æ–ø–µ—á–∞—Ç–∫–∞–º–∏\n- –ù–∞—Å—Ç—Ä–æ–∏—Ç—å –∞–≤—Ç–æ–º–∞—Ç–∏–∑–∏—Ä–æ–≤–∞–Ω–Ω—É—é –ø—Ä–æ–≤–µ—Ä–∫—É –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏ (–Ω–∞–ø—Ä–∏–º–µ—Ä, OWASP ZAP) –≤ CI/CD.")
    
    role_val = draft_data.get("role", "QA-–∏–Ω–∂–µ–Ω–µ—Ä")
    fullname_val = draft_data.get("fullname", "–ß–µ—Ä–∫–∞—Å–æ–≤ –ò–≥–æ—Ä—å")
    signature_date_val = draft_data.get("signature_date", "30.11.2025")
    
    # –û—á–∏—â–∞–µ–º session_state –ø–æ—Å–ª–µ –ø—Ä–∏–º–µ–Ω–µ–Ω–∏—è
    del st.session_state.draft_data
    del st.session_state.draft_modules
    del st.session_state.draft_defects
    del st.session_state.draft_saved_at
else:
    # –ó–Ω–∞—á–µ–Ω–∏—è –ø–æ —É–º–æ–ª—á–∞–Ω–∏—é
    report_title_val = "–û—Ç—á—ë—Ç –æ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏–∏ –º–æ–±–∏–ª—å–Ω–æ–≥–æ –ø—Ä–∏–ª–æ–∂–µ–Ω–∏—è –õ–µ–º–∞–Ω–∞ –ü–†–û"
    project_val = "–õ–µ–º–∞–Ω–∞ –ü–†–û"
    app_type_val = "–ú–æ–±–∏–ª—å–Ω–æ–µ"
    version_val = "241006.001"
    test_period_val = "29‚Äì30 –Ω–æ—è–±—Ä—è 2025 –≥."
    report_date_val = "30 –Ω–æ—è–±—Ä—è 2025 –≥."
    engineer_val = "–ß–µ—Ä–∫–∞—Å–æ–≤ –ò–≥–æ—Ä—å"
    
    release_status_val = "–ù–ï –†–ï–ö–û–ú–ï–ù–î–û–í–ê–ù –ö –í–´–ü–£–°–ö–£"
    s1_val = 2
    s2_val = 1
    total_tc_val = 72
    pass_tc_val = 69
    fail_tc_val = 3
    risk_val = "–£—è–∑–≤–∏–º–æ—Å—Ç–∏ –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏ –ø–æ–∑–≤–æ–ª—è—é—Ç –Ω–∞—Ä—É—à–∏—Ç–µ–ª—é –ø–æ–ª—É—á–∏—Ç—å –¥–æ—Å—Ç—É–ø –∫ –¥–∞–Ω–Ω—ã–º –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª–µ–π –∏ –≤—ã–∑–≤–∞—Ç—å –æ—Ç–∫–∞–∑ –≤ –æ–±—Å–ª—É–∂–∏–≤–∞–Ω–∏–∏."
    recommendation_val = "–†–µ–ª–∏–∑ –≤–æ–∑–º–æ–∂–µ–Ω —Ç–æ–ª—å–∫–æ –ø–æ—Å–ª–µ —É—Å—Ç—Ä–∞–Ω–µ–Ω–∏—è –≤—Å–µ—Ö S1/S2 –¥–µ—Ñ–µ–∫—Ç–æ–≤ –∏ –ø–æ–≤—Ç–æ—Ä–Ω–æ–≥–æ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è."
    
    device_browser_val = "Xiaomi 12"
    os_platform_val = "Android 15"
    build_val = "lemanna-pro_241006.001.apk"
    env_url_val = "https://test.lemanna.pro"
    tools_val = "Postman (API), Burp Suite (–±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç—å), Jira (–±–∞–≥-—Ç—Ä–µ–∫–∏–Ω–≥)"
    methodology_val = "–†—É—á–Ω–æ–µ —Ñ—É–Ω–∫—Ü–∏–æ–Ω–∞–ª—å–Ω–æ–µ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ + –ø—Ä–æ–≤–µ—Ä–∫–∞ –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏"
    
    consequences_val = "- S1 –¥–µ—Ñ–µ–∫—Ç—ã –ø–æ–∑–≤–æ–ª—è—é—Ç –∑–ª–æ—É–º—ã—à–ª–µ–Ω–Ω–∏–∫—É –ø–æ–ª—É—á–∏—Ç—å –¥–∞–Ω–Ω—ã–µ –¥—Ä—É–≥–∏—Ö –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª–µ–π –∏–ª–∏ –≤—ã–≤–µ—Å—Ç–∏ –ø—Ä–∏–ª–æ–∂–µ–Ω–∏–µ –∏–∑ —Å—Ç—Ä–æ—è.\n- S2 –¥–µ—Ñ–µ–∫—Ç —Å–Ω–∏–∂–∞–µ—Ç —é–∑–∞–±–∏–ª–∏—Ç–∏: –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª–∏ –Ω–µ –Ω–∞–π–¥—É—Ç —Ç–æ–≤–∞—Ä –ø—Ä–∏ –æ–ø–µ—á–∞—Ç–∫–µ."
    limitations_val = "1. –ù–µ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–ª–∞—Å—å –æ–ø–ª–∞—Ç–∞ —á–µ—Ä–µ–∑ Apple Pay (—É—Å—Ç—Ä–æ–π—Å—Ç–≤–æ Android).\n2. –ù–µ –ø—Ä–æ–≤–µ—Ä–µ–Ω–∞ —Å–∏–Ω—Ö—Ä–æ–Ω–∏–∑–∞—Ü–∏—è —Å 1–° (–Ω–µ—Ç –¥–æ—Å—Ç—É–ø–∞ –∫ –∏–Ω—Ç–µ–≥—Ä–∞—Ü–∏–æ–Ω–Ω–æ–º—É —Å—Ç–µ–Ω–¥—É).\n3. –ù–µ –ø—Ä–æ–≤–µ–¥–µ–Ω–æ –Ω–∞–≥—Ä—É–∑–æ—á–Ω–æ–µ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ (–æ–≥—Ä–∞–Ω–∏—á–µ–Ω–∏–µ –ø–æ –≤—Ä–µ–º–µ–Ω–∏)."
    conclusion_val = "–°–±–æ—Ä–∫–∞ 241006.001 —Å–æ–¥–µ—Ä–∂–∏—Ç –∫—Ä–∏—Ç–∏—á–µ—Å–∫–∏–µ —É—è–∑–≤–∏–º–æ—Å—Ç–∏ –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏, –¥–µ–ª–∞—é—â–∏–µ –µ—ë –Ω–µ–ø—Ä–∏–≥–æ–¥–Ω–æ–π –¥–ª—è –≤—ã–ø—É—Å–∫–∞ –≤ production. –ù–∞–ª–∏—á–∏–µ S1 –¥–µ—Ñ–µ–∫—Ç–æ–≤ –Ω–∞—Ä—É—à–∞–µ—Ç –±–∞–∑–æ–≤—ã–µ –ø—Ä–∏–Ω—Ü–∏–ø—ã –∑–∞—â–∏—Ç—ã –¥–∞–Ω–Ω—ã—Ö –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª–µ–π."
    recommendations_detailed_val = "–ù–µ–º–µ–¥–ª–µ–Ω–Ω–æ –∏—Å–ø—Ä–∞–≤–∏—Ç—å —É—è–∑–≤–∏–º–æ—Å—Ç–∏ BUG-SEC-001 –∏ BUG-SEC-002.\n–†–µ–∞–ª–∏–∑–æ–≤–∞—Ç—å fuzzy search –¥–ª—è –ø–æ–≤—ã—à–µ–Ω–∏—è —é–∑–∞–±–∏–ª–∏—Ç–∏ (BUG-SEARCH-001).\n–ü—Ä–æ–≤–µ—Å—Ç–∏ –ø–æ–≤—Ç–æ—Ä–Ω–æ–µ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ –ø–æ—Å–ª–µ —Ñ–∏–∫—Å–æ–≤ —Å —Ñ–æ–∫—É—Å–æ–º –Ω–∞:\n- –ü–æ–≤—Ç–æ—Ä–Ω—É—é –ø—Ä–æ–≤–µ—Ä–∫—É –ø–æ–ª–µ–π –≤–≤–æ–¥–∞ –Ω–∞ –∏–Ω—ä–µ–∫—Ü–∏–∏\n- –¢–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ —Å—Ü–µ–Ω–∞—Ä–∏–µ–≤ –ø–æ–∏—Å–∫–∞ —Å –æ–ø–µ—á–∞—Ç–∫–∞–º–∏\n- –ù–∞—Å—Ç—Ä–æ–∏—Ç—å –∞–≤—Ç–æ–º–∞—Ç–∏–∑–∏—Ä–æ–≤–∞–Ω–Ω—É—é –ø—Ä–æ–≤–µ—Ä–∫—É –±–µ–∑–æ–ø–∞—Å–Ω–æ—Å—Ç–∏ (–Ω–∞–ø—Ä–∏–º–µ—Ä, OWASP ZAP) –≤ CI/CD."
    
    role_val = "QA-–∏–Ω–∂–µ–Ω–µ—Ä"
    fullname_val = "–ß–µ—Ä–∫–∞—Å–æ–≤ –ò–≥–æ—Ä—å"
    signature_date_val = "30.11.2025"

# === –§–û–†–ú–ê –í–í–û–î–ê –î–ê–ù–ù–´–• ===
with st.form("main_form"):
    # –ó–∞–≥–æ–ª–æ–≤–æ–∫ –∏ –æ—Å–Ω–æ–≤–Ω–∞—è –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è
    st.subheader("1. –ó–∞–≥–æ–ª–æ–≤–æ–∫ –∏ –æ—Å–Ω–æ–≤–Ω–∞—è –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è")
    
    report_title = st.text_input("–ù–∞–∑–≤–∞–Ω–∏–µ –æ—Ç—á—ë—Ç–∞", report_title_val)
    col1, col2, col3 = st.columns(3)
    with col1:
        project = st.text_input("–ü—Ä–æ–µ–∫—Ç", project_val)
        app_type = st.selectbox("–¢–∏–ø –ø—Ä–∏–ª–æ–∂–µ–Ω–∏—è", ["–ú–æ–±–∏–ª—å–Ω–æ–µ", "–í–µ–±"], 
                               index=0 if app_type_val == "–ú–æ–±–∏–ª—å–Ω–æ–µ" else 1)
    with col2:
        version = st.text_input("–í–µ—Ä—Å–∏—è", version_val)
        test_period = st.text_input("–ü–µ—Ä–∏–æ–¥ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è", test_period_val)
    with col3:
        report_date = st.text_input("–î–∞—Ç–∞ —Ñ–æ—Ä–º–∏—Ä–æ–≤–∞–Ω–∏—è –æ—Ç—á—ë—Ç–∞", report_date_val)
        engineer = st.text_input("–ò–Ω–∂–µ–Ω–µ—Ä –ø–æ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—é", engineer_val)
    
    release_status = st.selectbox(
        "–°—Ç–∞—Ç—É—Å —Ä–µ–ª–∏–∑–∞",
        ["–ù–ï –†–ï–ö–û–ú–ï–ù–î–û–í–ê–ù –ö –í–´–ü–£–°–ö–£", "–†–ï–ö–û–ú–ï–ù–î–û–í–ê–ù –ö –í–´–ü–£–°–ö–£ –° –ó–ê–ú–ï–ß–ê–ù–ò–Ø–ú–ò", "–†–ï–ö–û–ú–ï–ù–î–û–í–ê–ù –ö –í–´–ü–£–°–ö–£"],
        index=["–ù–ï –†–ï–ö–û–ú–ï–ù–î–û–í–ê–ù –ö –í–´–ü–£–°–ö–£", "–†–ï–ö–û–ú–ï–ù–î–û–í–ê–ù –ö –í–´–ü–£–°–ö–£ –° –ó–ê–ú–ï–ß–ê–ù–ò–Ø–ú–ò", "–†–ï–ö–û–ú–ï–ù–î–û–í–ê–ù –ö –í–´–ü–£–°–ö–£"].index(release_status_val)
    )
    
    # –ú–µ—Ç—Ä–∏–∫–∏
    st.subheader("2. –ú–µ—Ç—Ä–∏–∫–∏ —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        s1 = st.number_input("–î–µ—Ñ–µ–∫—Ç—ã S1", min_value=0, value=s1_val)
    with col2:
        s2 = st.number_input("–î–µ—Ñ–µ–∫—Ç—ã S2", min_value=0, value=s2_val)
    with col3:
        total_tc = st.number_input("–í—Å–µ–≥–æ —Ç–µ—Å—Ç-–∫–µ–π—Å–æ–≤", min_value=0, value=total_tc_val)
    with col4:
        pass_tc = st.number_input("–ü—Ä–æ–π–¥–µ–Ω–æ —É—Å–ø–µ—à–Ω–æ (PASS)", min_value=0, value=pass_tc_val)
    fail_tc = st.number_input("–£–ø–∞–ª–æ (FAIL)", min_value=0, value=fail_tc_val)
    
    # –†–∏—Å–∫–∏ –∏ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏
    risk = st.text_area("–†–∏—Å–∫–∏", risk_val, height=80)
    recommendation = st.text_area("–ö—Ä–∞—Ç–∫–∏–µ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏", recommendation_val, height=60)
    
    # –ö–æ–Ω—Ç–µ–∫—Å—Ç —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è
    st.subheader("3. –ö–æ–Ω—Ç–µ–∫—Å—Ç —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è")
    col1, col2 = st.columns(2)
    with col1:
        device_browser = st.text_input("–£—Å—Ç—Ä–æ–π—Å—Ç–≤–æ / –ë—Ä–∞—É–∑–µ—Ä", device_browser_val)
        os_platform = st.text_input("–û–° / –ü–ª–∞—Ç—Ñ–æ—Ä–º–∞", os_platform_val)
        build = st.text_input("–°–±–æ—Ä–∫–∞", build_val)
        env_url = st.text_input("URL –æ–∫—Ä—É–∂–µ–Ω–∏—è", env_url_val)
    with col2:
        tools = st.text_area("–ò–Ω—Å—Ç—Ä—É–º–µ–Ω—Ç—ã", tools_val, height=100)
        methodology = st.text_area("–ú–µ—Ç–æ–¥–æ–ª–æ–≥–∏—è —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è", methodology_val, height=100)
    
    # –†–µ–∑—É–ª—å—Ç–∞—Ç—ã –ø–æ –º–æ–¥—É–ª—è–º
    st.subheader("4. –†–µ–∑—É–ª—å—Ç–∞—Ç—ã –ø–æ –º–æ–¥—É–ª—è–º")
    module_data_list = []
    
    for i, mod in enumerate(default_modules):
        with st.expander(f"–ú–æ–¥—É–ª—å {i+1}: {mod['title']}", expanded=False):
            title = st.text_input(f"–ù–∞–∑–≤–∞–Ω–∏–µ –º–æ–¥—É–ª—è {i+1}", mod["title"], key=f"title_{i}")
            df_edited = st.data_editor(
                mod["df"],
                column_config={
                    "ID": st.column_config.TextColumn("ID", width="small"),
                    "–°—Ü–µ–Ω–∞—Ä–∏–π": st.column_config.TextColumn("–°—Ü–µ–Ω–∞—Ä–∏–π", width="medium"),
                    "–°—Ç–∞—Ç—É—Å": st.column_config.SelectboxColumn("–°—Ç–∞—Ç—É—Å", options=["PASS", "FAIL", "SKIP"], width="small"),
                    "–ö–æ–º–º–µ–Ω—Ç–∞—Ä–∏–π": st.column_config.TextColumn("–ö–æ–º–º–µ–Ω—Ç–∞—Ä–∏–π", width="large"),
                },
                hide_index=True,
                key=f"module_{i}",
                use_container_width=True,
                num_rows="dynamic"
            )
            module_data_list.append({"title": title, "df": df_edited})
    
    # –ê–Ω–∞–ª–∏–∑ –¥–µ—Ñ–µ–∫—Ç–æ–≤
    st.subheader("5. –ê–Ω–∞–ª–∏–∑ –¥–µ—Ñ–µ–∫—Ç–æ–≤")
    defects = st.data_editor(
        default_defects if "draft_defects" not in st.session_state else draft_defects,
        column_config={
            "ID": st.column_config.TextColumn("ID", width="small"),
            "–ú–æ–¥—É–ª—å": st.column_config.TextColumn("–ú–æ–¥—É–ª—å", width="small"),
            "–ó–∞–≥–æ–ª–æ–≤–æ–∫": st.column_config.TextColumn("–ó–∞–≥–æ–ª–æ–≤–æ–∫", width="medium"),
            "–°–µ—Ä—å—ë–∑–Ω–æ—Å—Ç—å": st.column_config.SelectboxColumn("–°–µ—Ä—å—ë–∑–Ω–æ—Å—Ç—å", options=["S1", "S2", "S3", "S4"], width="small"),
            "–°—Ç–∞—Ç—É—Å": st.column_config.TextColumn("–°—Ç–∞—Ç—É—Å", width="small"),
        },
        hide_index=True,
        use_container_width=True,
        num_rows="dynamic"
    )
    
    # –û–≥—Ä–∞–Ω–∏—á–µ–Ω–∏—è –∏ –≤—ã–≤–æ–¥
    st.subheader("6. –û–≥—Ä–∞–Ω–∏—á–µ–Ω–∏—è –∏ –≤—ã–≤–æ–¥")
    limitations = st.text_area(
        "–û–≥—Ä–∞–Ω–∏—á–µ–Ω–∏—è —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è (–Ω—É–º–µ—Ä–æ–≤–∞–Ω–Ω—ã–π —Å–ø–∏—Å–æ–∫)",
        limitations_val,
        height=120,
        help="–í–≤–µ–¥–∏—Ç–µ –ø–æ –æ–¥–Ω–æ–º—É –æ–≥—Ä–∞–Ω–∏—á–µ–Ω–∏—é –Ω–∞ —Å—Ç—Ä–æ–∫—É. –ú–∞—Ä–∫–µ—Ä—ã –±—É–¥—É—Ç –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏ –ø—Ä–µ–æ–±—Ä–∞–∑–æ–≤–∞–Ω—ã –≤ –Ω—É–º–µ—Ä–∞—Ü–∏—é."
    )
    consequences = st.text_area("–ü–æ—Å–ª–µ–¥—Å—Ç–≤–∏—è –¥–µ—Ñ–µ–∫—Ç–æ–≤", consequences_val, height=100)
    conclusion = st.text_area("–í—ã–≤–æ–¥", conclusion_val, height=100)
    recommendations_detailed = st.text_area("–î–µ—Ç–∞–ª—å–Ω—ã–µ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏", recommendations_detailed_val, height=100)
    
    # –ü–æ–¥–ø–∏—Å—å
    st.subheader("7. –ü–æ–¥–ø–∏—Å—å")
    col1, col2, col3 = st.columns(3)
    with col1:
        role = st.text_input("–†–æ–ª—å", role_val)
    with col2:
        fullname = st.text_input("–§–ò–û", fullname_val)
    with col3:
        signature_date = st.text_input("–î–∞—Ç–∞", signature_date_val)
    
    # === –ö–ù–û–ü–ö–ò –§–û–†–ú–´ ===
    col1, col2 = st.columns(2)
    with col1:
        save_draft_clicked = st.form_submit_button("üíæ –°–æ—Ö—Ä–∞–Ω–∏—Ç—å —á–µ—Ä–Ω–æ–≤–∏–∫", type="secondary")
    with col2:
        submitted = st.form_submit_button("üì• –°–æ–∑–¥–∞—Ç—å –æ—Ç—á—ë—Ç", type="primary")

# === –û–ë–†–ê–ë–û–¢–ö–ê –ö–ù–û–ü–û–ö (–í–ù–ï –§–û–†–ú–´!) ===
if save_draft_clicked:
    # –°–æ–±–∏—Ä–∞–µ–º –¥–∞–Ω–Ω—ã–µ –∏–∑ —Ç–µ–∫—É—â–µ–π —Ñ–æ—Ä–º—ã
    data = {
        "report_title": report_title,
        "project": project,
        "app_type": app_type,
        "version": version,
        "test_period": test_period,
        "report_date": report_date,
        "engineer": engineer,
        "release_status": release_status,
        "s1": s1,
        "s2": s2,
        "total_tc": total_tc,
        "pass": pass_tc,
        "fail": fail_tc,
        "device_browser": device_browser,
        "os_platform": os_platform,
        "build": build,
        "env_url": env_url.strip(),
        "tools": tools,
        "methodology": methodology,
        "risk": risk,
        "recommendation": recommendation,
        "consequences": consequences,
        "limitations": limitations,
        "conclusion": conclusion,
        "recommendations_detailed": recommendations_detailed,
        "role": role,
        "fullname": fullname,
        "signature_date": signature_date,
    }
    
    draft_json = save_draft(data, module_data_list, defects)
    st.session_state.draft_to_download = draft_json
    st.session_state.draft_filename = f"—á–µ—Ä–Ω–æ–≤–∏–∫_{datetime.now().strftime('%d%m%Y_%H%M%S')}.json"
    st.success("‚úÖ –ß–µ—Ä–Ω–æ–≤–∏–∫ –ø–æ–¥–≥–æ—Ç–æ–≤–ª–µ–Ω –∫ —Å–∫–∞—á–∏–≤–∞–Ω–∏—é!")
    st.rerun()

# –ö–Ω–æ–ø–∫–∞ —Å–∫–∞—á–∏–≤–∞–Ω–∏—è —á–µ—Ä–Ω–æ–≤–∏–∫–∞ (–≤–Ω–µ —Ñ–æ—Ä–º—ã)
if "draft_to_download" in st.session_state and st.session_state.draft_to_download is not None:
    st.download_button(
        "‚¨áÔ∏è –°–∫–∞—á–∞—Ç—å —á–µ—Ä–Ω–æ–≤–∏–∫",
        st.session_state.draft_to_download,
        st.session_state.draft_filename,
        "application/json",
        use_container_width=True,
        type="primary"
    )
    if st.button("–ó–∞–∫—Ä—ã—Ç—å", key="close_draft"):
        del st.session_state.draft_to_download
        del st.session_state.draft_filename
        st.rerun()

# === –ì–ï–ù–ï–†–ê–¶–ò–Ø –û–¢–ß–Å–¢–ê ===
if submitted:
    # –°–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö
    data = {
        "report_title": report_title,
        "project": project,
        "app_type": app_type,
        "version": version,
        "test_period": test_period,
        "report_date": report_date,
        "engineer": engineer,
        "release_status": release_status,
        "s1": s1,
        "s2": s2,
        "total_tc": total_tc,
        "pass": pass_tc,
        "fail": fail_tc,
        "device_browser": device_browser,
        "os_platform": os_platform,
        "build": build,
        "env_url": env_url.strip(),
        "tools": tools,
        "methodology": methodology,
        "risk": risk,
        "recommendation": recommendation,
        "consequences": consequences,
        "limitations": limitations,
        "conclusion": conclusion,
        "recommendations_detailed": recommendations_detailed,
        "role": role,
        "fullname": fullname,
        "signature_date": signature_date,
    }
    
    # –í–∞–ª–∏–¥–∞—Ü–∏—è
    validation_errors = []
    if pass_tc + fail_tc != total_tc:
        validation_errors.append("‚ö†Ô∏è –°—É–º–º–∞ —Å—Ç–∞—Ç—É—Å–æ–≤ (PASS + FAIL) –Ω–µ —Ä–∞–≤–Ω–∞ –æ–±—â–µ–º—É –∫–æ–ª–∏—á–µ—Å—Ç–≤—É —Ç–µ—Å—Ç-–∫–µ–π—Å–æ–≤")
    
    if validation_errors:
        for err in validation_errors:
            st.error(err)
        st.stop()
    
    # –ì–µ–Ω–µ—Ä–∞—Ü–∏—è –æ—Ç—á—ë—Ç–æ–≤
    with st.spinner("–ì–µ–Ω–µ—Ä–∞—Ü–∏—è –æ—Ç—á—ë—Ç–æ–≤..."):
        # DOCX
        docx_buffer = generate_docx(data, module_data_list, defects)
        
        # HTML
        html_content = generate_html_report(data, module_data_list, defects)
        html_buffer = io.BytesIO(html_content.encode('utf-8'))
        
        # XLSX
        xlsx_buffer = generate_xlsx_single_sheet(data, module_data_list, defects)
    
    # –û—Ç–æ–±—Ä–∞–∂–µ–Ω–∏–µ –∫–Ω–æ–ø–æ–∫ –∑–∞–≥—Ä—É–∑–∫–∏
    st.success("‚úÖ –û—Ç—á—ë—Ç—ã —É—Å–ø–µ—à–Ω–æ —Å–≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞–Ω—ã!")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.download_button(
            "üìÑ –°–∫–∞—á–∞—Ç—å DOCX",
            docx_buffer,
            "–æ—Ç—á—ë—Ç_—Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    
    with col2:
        st.download_button(
            "üåê –°–∫–∞—á–∞—Ç—å HTML",
            html_buffer,
            "–æ—Ç—á—ë—Ç_—Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ.html",
            "text/html",
            use_container_width=True
        )
    
    with col3:
        st.download_button(
            "üìä –°–∫–∞—á–∞—Ç—å XLSX",
            xlsx_buffer,
            "–æ—Ç—á—ë—Ç_—Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )
    
    # –ü—Ä–µ–¥–ø—Ä–æ—Å–º–æ—Ç—Ä HTML
    with st.expander("üîç –ü—Ä–µ–¥–ø—Ä–æ—Å–º–æ—Ç—Ä HTML-–æ—Ç—á—ë—Ç–∞"):
        st.components.v1.html(html_content, height=600, scrolling=True)